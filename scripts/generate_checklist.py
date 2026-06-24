import os
import sys

# Define brand parameters
BRAND_NAME = "Intikhab"
DOC_TITLE = "INTIKHAB WEBSITE CLIENT TESTING CHECKLIST"
DOC_SUBTITLE = "Pre-Launch Review & Approval Checklist"
INTRO_TEXT = (
    "Thank you for reviewing the Intikhab platform. Please go through the following checklist and mark "
    "each item under the corresponding status column (Passed, Needs Changes, or Not Tested). "
    "Include comments or required adjustments in the spaces provided below each section."
)

SECTIONS = [
    {
        "title": "HOMEPAGE",
        "items": [
            "Homepage loads correctly",
            "Logo displays correctly",
            "Navigation menu works",
            "Search bar works",
            "Featured products display correctly",
            "Categories display correctly",
            "Mobile homepage looks good",
            "No spelling mistakes",
            "Images load correctly"
        ]
    },
    {
        "title": "PRODUCT BROWSING",
        "items": [
            "Product listing page loads",
            "Filters work correctly",
            "Search returns correct results",
            "Product cards display correctly",
            "Product images load correctly",
            "Sale badges display correctly",
            "Pagination/load more works",
            "Category filtering works"
        ]
    },
    {
        "title": "PRODUCT DETAIL PAGE",
        "items": [
            "Product page opens correctly",
            "Product images gallery works",
            "Product information is correct",
            "Sizes display correctly",
            "Stock status is correct",
            "Related products display correctly",
            "Reviews section works"
        ]
    },
    {
        "title": "CART",
        "items": [
            "Add to cart works",
            "Remove item works",
            "Update quantity works",
            "Multiple sizes remain separate",
            "Cart totals update correctly",
            "Cart drawer works",
            "Cart page works"
        ]
    },
    {
        "title": "CHECKOUT",
        "items": [
            "Checkout page loads",
            "Customer information can be entered",
            "Address fields work",
            "Postal code field works",
            "Coupon field works",
            "Totals calculate correctly",
            "Order can be submitted",
            "Order confirmation appears"
        ]
    },
    {
        "title": "USER ACCOUNT",
        "items": [
            "Sign up works",
            "Login works",
            "Logout works",
            "Account page works",
            "Order history displays",
            "Saved addresses work",
            "Wishlist works"
        ]
    },
    {
        "title": "ADMIN PANEL",
        "items": [
            "Admin login works",
            "Dashboard loads",
            "Products management works",
            "Orders management works",
            "Categories management works",
            "Reviews moderation works",
            "Coupons management works",
            "Feedback/messages display correctly"
        ]
    },
    {
        "title": "MOBILE TESTING",
        "items": [
            "Homepage works on mobile",
            "Product pages work on mobile",
            "Cart works on mobile",
            "Checkout works on mobile",
            "Navigation works on mobile",
            "No layout issues",
            "No text overlaps"
        ]
    },
    {
        "title": "PERFORMANCE & QUALITY",
        "items": [
            "Website feels fast",
            "No broken pages",
            "No broken images",
            "No visual glitches",
            "No console errors noticed",
            "Overall design quality approved"
        ]
    }
]

def generate_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    pdf_filename = "Intikhab_QA_Testing_Checklist.pdf"
    print(f"Generating PDF: {pdf_filename}...")

    # Document Layout: Letter size, 0.5-inch margins for print efficiency
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    # Premium Modern Theme Colors (Matching Intikhab Web Theme)
    c_charcoal = colors.HexColor('#1D1A17') # Brand Charcoal
    c_crimson = colors.HexColor('#B42318')  # Brand Crimson
    c_border = colors.HexColor('#E4D8C8')   # Brand Sand Border
    c_bg_light = colors.HexColor('#F6EFE5') # Column Header Background
    c_surface = colors.HexColor('#FFFDF8')  # Secondary Shading

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=c_charcoal,
        spaceAfter=3
    )
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=c_crimson,
        spaceAfter=10
    )
    intro_style = ParagraphStyle(
        'DocIntro',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#4B5563'),
        spaceAfter=12
    )
    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=13,
        textColor=colors.white,
    )
    col_header_left_style = ParagraphStyle(
        'ColHeaderLeft',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=c_charcoal
    )
    col_header_center_style = ParagraphStyle(
        'ColHeaderCenter',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        alignment=1, # Center
        textColor=c_charcoal
    )
    item_style = ParagraphStyle(
        'ItemStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor('#111827')
    )
    center_checkbox_style = ParagraphStyle(
        'CenterCheckbox',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=10.5,
        alignment=1, # Center
        textColor=colors.HexColor('#9CA3AF') # Gray bracket un-checked
    )
    comment_box_label_style = ParagraphStyle(
        'CommentLabel',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=9,
        textColor=colors.HexColor('#4B5563')
    )

    story = []

    # 1. Page Header (Modern Executive Layout)
    story.append(Paragraph(DOC_TITLE, title_style))
    story.append(Paragraph(DOC_SUBTITLE, subtitle_style))
    story.append(Paragraph(INTRO_TEXT, intro_style))
    story.append(Spacer(1, 6))

    # Column Widths matching 540pt total printable area (612pt letter page width - 72pt margins)
    col_widths = [330, 65, 85, 60]

    # 2. Sections
    for idx, section in enumerate(SECTIONS, 1):
        section_story = []
        
        # Build checklist table data
        # Row 0: Section header title span
        header_para = Paragraph(f"SECTION {idx} — {section['title'].upper()}", section_title_style)
        
        # Row 1: Column subheaders
        col_item = Paragraph("Checklist Item", col_header_left_style)
        col_passed = Paragraph("Passed", col_header_center_style)
        col_needs = Paragraph("Needs Changes", col_header_center_style)
        col_not = Paragraph("Not Tested", col_header_center_style)

        table_data = [
            [header_para, "", "", ""],
            [col_item, col_passed, col_needs, col_not]
        ]

        # Add checklist items
        for item in section["items"]:
            desc = Paragraph(item, item_style)
            chk_passed = Paragraph("[ &nbsp; ]", center_checkbox_style)
            chk_needs = Paragraph("[ &nbsp; ]", center_checkbox_style)
            chk_not = Paragraph("[ &nbsp; ]", center_checkbox_style)
            table_data.append([desc, chk_passed, chk_needs, chk_not])

        checklist_table = Table(table_data, colWidths=col_widths)
        
        # Apply structured table styles
        t_style = TableStyle([
            # Section Title Row
            ('SPAN', (0,0), (3,0)),
            ('BACKGROUND', (0,0), (3,0), c_charcoal),
            ('TOPPADDING', (0,0), (3,0), 5),
            ('BOTTOMPADDING', (0,0), (3,0), 5),
            ('LEFTPADDING', (0,0), (3,0), 6),
            
            # Column Subheaders Row
            ('BACKGROUND', (0,1), (3,1), c_bg_light),
            ('TOPPADDING', (0,1), (3,1), 4),
            ('BOTTOMPADDING', (0,1), (3,1), 4),
            ('LEFTPADDING', (0,1), (3,1), 6),
            ('LINEBELOW', (0,1), (3,1), 1.25, c_charcoal),
            
            # Global properties
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,2), (0,-1), 6),
            ('LINEBELOW', (0,2), (-1,-1), 0.5, c_border),
        ])

        # Add alternating row backgrounds starting from row 2
        for r_idx in range(2, len(table_data)):
            bg_color = c_surface if r_idx % 2 == 0 else colors.white
            t_style.add('BACKGROUND', (0, r_idx), (-1, r_idx), bg_color)

        checklist_table.setStyle(t_style)
        section_story.append(checklist_table)
        section_story.append(Spacer(1, 3))

        # Comments Box
        comm_label = Paragraph("COMMENTS / REQUIRED ADJUSTMENTS:", comment_box_label_style)
        comments_box = Table([[comm_label]], colWidths=[540], rowHeights=[42])
        comments_box.setStyle(TableStyle([
            ('BOX', (0,0), (-1,-1), 0.5, c_border),
            ('BACKGROUND', (0,0), (-1,-1), c_surface),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
        ]))
        section_story.append(comments_box)
        section_story.append(Spacer(1, 12))

        # Keep the section table and its comment box together
        story.append(KeepTogether(section_story))

    # 3. Section 10: Final Approval
    approval_story = []
    
    header_para = Paragraph("SECTION 10 — FINAL APPROVAL", section_title_style)
    header_table = Table([[header_para, "", "", ""]], colWidths=col_widths)
    header_table.setStyle(TableStyle([
        ('SPAN', (0,0), (3,0)),
        ('BACKGROUND', (0,0), (3,0), c_charcoal),
        ('TOPPADDING', (0,0), (3,0), 5),
        ('BOTTOMPADDING', (0,0), (3,0), 5),
        ('LEFTPADDING', (0,0), (3,0), 6),
    ]))
    approval_story.append(header_table)
    approval_story.append(Spacer(1, 4))

    rating_desc = Paragraph("<b>Overall Rating:</b>", item_style)
    rating_options = Paragraph("[ &nbsp; ] Excellent &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; [ &nbsp; ] Good &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; [ &nbsp; ] Needs Minor Changes &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; [ &nbsp; ] Needs Major Changes", item_style)
    
    launch_desc = Paragraph("<b>Would you approve this version for launch?</b>", item_style)
    launch_options = Paragraph("[ &nbsp; ] Yes &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; [ &nbsp; ] No", item_style)

    approval_table_data = [
        [rating_desc, rating_options],
        [launch_desc, launch_options]
    ]
    approval_table = Table(approval_table_data, colWidths=[200, 340])
    approval_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, c_border),
    ]))
    approval_story.append(approval_table)
    approval_story.append(Spacer(1, 4))

    feedback_label = Paragraph("ADDITIONAL CLIENT FEEDBACK:", comment_box_label_style)
    feedback_box = Table([[feedback_label]], colWidths=[540], rowHeights=[60])
    feedback_box.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, c_border),
        ('BACKGROUND', (0,0), (-1,-1), c_surface),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
    ]))
    approval_story.append(feedback_box)
    approval_story.append(Spacer(1, 6))

    # Signature fields table
    sig_label_client = Paragraph("Client Name (Printed): _________________________________________", item_style)
    sig_label_date = Paragraph("Date: _________________________", item_style)
    sig_label_sig = Paragraph("Signature: __________________________________________________", item_style)

    sig_table_data = [
        [sig_label_client, sig_label_date],
        [sig_label_sig, Paragraph("", item_style)]
    ]
    sig_table = Table(sig_table_data, colWidths=[340, 200])
    sig_table.setStyle(TableStyle([
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    approval_story.append(sig_table)
    story.append(KeepTogether(approval_story))

    # 4. Page Break & Final Page Visual Summary Dashboard
    story.append(PageBreak())
    
    summary_story = []
    summary_title = Paragraph("INTIKHAB PLATFORM VERIFICATION SUMMARY", title_style)
    summary_subtitle = Paragraph("Launch Readiness Dashboard Status", subtitle_style)
    summary_story.append(summary_title)
    summary_story.append(summary_subtitle)
    summary_story.append(Spacer(1, 6))

    sum_col_item = Paragraph("Platform Module Component", col_header_left_style)
    sum_col_passed = Paragraph("Passed", col_header_center_style)
    sum_col_needs = Paragraph("Needs Changes", col_header_center_style)
    sum_col_not = Paragraph("Not Tested", col_header_center_style)
    
    summary_rows = [
        [sum_col_item, sum_col_passed, sum_col_needs, sum_col_not]
    ]

    modules = [
        "Homepage",
        "Products",
        "Cart",
        "Checkout",
        "Account",
        "Admin",
        "Mobile",
        "Performance"
    ]

    for mod in modules:
        col1 = Paragraph(f"<b>{mod}</b>", item_style)
        col2 = Paragraph("[ &nbsp; ]", center_checkbox_style)
        col3 = Paragraph("[ &nbsp; ]", center_checkbox_style)
        col4 = Paragraph("[ &nbsp; ]", center_checkbox_style)
        summary_rows.append([col1, col2, col3, col4])

    summary_table = Table(summary_rows, colWidths=col_widths)
    
    # Dashboard styling
    sum_t_style = TableStyle([
        # Headers
        ('BACKGROUND', (0,0), (3,0), c_bg_light),
        ('TOPPADDING', (0,0), (3,0), 6),
        ('BOTTOMPADDING', (0,0), (3,0), 6),
        ('LEFTPADDING', (0,0), (3,0), 6),
        ('LINEABOVE', (0,0), (3,0), 1.25, c_charcoal),
        ('LINEBELOW', (0,0), (3,0), 1.25, c_charcoal),
        
        # Grid lines
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,1), (0,-1), 6),
        ('LINEBELOW', (0,1), (-1,-1), 0.5, c_border),
    ])

    for r_idx in range(1, len(summary_rows)):
        bg_color = c_surface if r_idx % 2 == 1 else colors.white
        sum_t_style.add('BACKGROUND', (0, r_idx), (-1, r_idx), bg_color)

    summary_table.setStyle(sum_t_style)
    summary_story.append(summary_table)
    summary_story.append(Spacer(1, 20))

    footer_text = Paragraph(
        "<i>Intikhab E-Commerce Platform Client Checklist — Confidential Review Copy. "
        "Generated via automated testing suites.</i>",
        comment_box_label_style
    )
    summary_story.append(footer_text)
    
    # Crucial Fix: Use extend instead of append to prevent appending a list directly into story
    story.extend(summary_story)

    # Build PDF
    doc.build(story)
    print("PDF generated successfully.")

def generate_docx():
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    docx_filename = "Intikhab_QA_Testing_Checklist.docx"
    print(f"Generating DOCX: {docx_filename}...")

    doc = docx.Document()

    # Set Margins (0.5 inch margins for compactness to match layout scaling)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    # Styling helper for tables
    def style_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # XML Helper to align text vertically to center in table cell
    def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>\n'
            f'  <w:top w:w="{top}" w:type="dxa"/>\n'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
            f'  <w:left w:w="{left}" w:type="dxa"/>\n'
            f'  <w:right w:w="{right}" w:type="dxa"/>\n'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(DOC_TITLE)
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(20)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(29, 26, 23) # Charcoal
    p_title.paragraph_format.space_after = Pt(2)

    # Subtitle
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run(DOC_SUBTITLE)
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(11)
    r_sub.bold = True
    r_sub.font.color.rgb = RGBColor(180, 35, 24) # Crimson Red
    p_sub.paragraph_format.space_after = Pt(8)

    # Intro
    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run(INTRO_TEXT)
    r_intro.font.name = 'Arial'
    r_intro.font.size = Pt(9)
    r_intro.font.color.rgb = RGBColor(75, 85, 99)
    p_intro.paragraph_format.space_after = Pt(14)

    # Sections
    for idx, section in enumerate(SECTIONS, 1):
        # Create Table with 4 columns
        # Total printable width is 7.5 inches
        col_widths = [Inches(4.3), Inches(1.0), Inches(1.2), Inches(1.0)]
        
        table = doc.add_table(rows=0, cols=4)
        table.autofit = False
        style_table_borders(table)

        # 1. Section Header Row (Merged all 4 cols)
        row_header = table.add_row()
        cell_header = row_header.cells[0].merge(row_header.cells[1]).merge(row_header.cells[2]).merge(row_header.cells[3])
        cell_header.width = Inches(7.5)
        set_cell_background(cell_header, "1D1A17")
        set_cell_margins(cell_header, top=100, bottom=100, left=120, right=120)
        
        hp = cell_header.paragraphs[0]
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hrun = hp.add_run(f"SECTION {idx} — {section['title'].upper()}")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(10)
        hrun.bold = True
        hrun.font.color.rgb = RGBColor(255, 255, 255)

        # 2. Subheader columns row
        row_sub = table.add_row()
        sub_headers = ["Checklist Item", "Passed", "Needs Changes", "Not Tested"]
        for c_idx, sub_h in enumerate(sub_headers):
            cell = row_sub.cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, "F6EFE5")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(sub_h)
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.bold = True
            run.font.color.rgb = RGBColor(29, 26, 23)

        # 3. Add Checklist Items
        for r_offset, item in enumerate(section["items"]):
            row = table.add_row()
            bg_color = "FFFDF8" if r_offset % 2 == 0 else "FFFFFF"
            
            # Setup checklist items in each cell
            for c_idx in range(4):
                cell = row.cells[c_idx]
                cell.width = col_widths[c_idx]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)

                if c_idx == 0:
                    run = p.add_run(item)
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(17, 24, 39)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("[   ]")
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(156, 163, 175)

        # Comments Box right below the table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(3)
        p_space.paragraph_format.space_after = Pt(3)

        comm_table = doc.add_table(rows=1, cols=1)
        comm_table.autofit = False
        comm_cell = comm_table.cell(0, 0)
        set_cell_background(comm_cell, "FFFDF8")
        comm_cell.width = Inches(7.5)
        set_cell_margins(comm_cell, top=80, bottom=500, left=100, right=100) # Extra bottom margin to mimic height
        
        tcPr = comm_cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

        cp = comm_cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        crun = cp.add_run("COMMENTS / REQUIRED ADJUSTMENTS:")
        crun.font.name = 'Arial'
        crun.font.size = Pt(7.5)
        crun.bold = True
        crun.font.color.rgb = RGBColor(75, 85, 99)

        # Add spacing before next section
        p_sec_space = doc.add_paragraph()
        p_sec_space.paragraph_format.space_after = Pt(10)

    # Section 10: Final Approval
    app_table = doc.add_table(rows=0, cols=4)
    app_table.autofit = False
    style_table_borders(app_table)

    col_widths = [Inches(4.3), Inches(1.0), Inches(1.2), Inches(1.0)]

    row_header = app_table.add_row()
    cell_header = row_header.cells[0].merge(row_header.cells[1]).merge(row_header.cells[2]).merge(row_header.cells[3])
    cell_header.width = Inches(7.5)
    set_cell_background(cell_header, "1D1A17")
    set_cell_margins(cell_header, top=100, bottom=100, left=120, right=120)
    
    app_p = cell_header.paragraphs[0]
    app_p.paragraph_format.space_after = Pt(0)
    app_run = app_p.add_run("SECTION 10 — FINAL APPROVAL")
    app_run.font.name = 'Arial'
    app_run.font.size = Pt(10)
    app_run.bold = True
    app_run.font.color.rgb = RGBColor(255, 255, 255)

    p_spacer_ap = doc.add_paragraph()
    p_spacer_ap.paragraph_format.space_before = Pt(3)
    p_spacer_ap.paragraph_format.space_after = Pt(3)

    # Options Table
    opt_table = doc.add_table(rows=2, cols=2)
    opt_table.autofit = False
    style_table_borders(opt_table)
    
    opt_table.rows[0].cells[0].width = Inches(2.8)
    opt_table.rows[0].cells[1].width = Inches(4.7)
    opt_table.rows[1].cells[0].width = Inches(2.8)
    opt_table.rows[1].cells[1].width = Inches(4.7)

    for r in opt_table.rows:
        for c in r.cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    p_opt0 = opt_table.rows[0].cells[0].paragraphs[0]
    p_opt0.paragraph_format.space_after = Pt(0)
    r_opt0 = p_opt0.add_run("Overall Rating:")
    r_opt0.bold = True
    r_opt0.font.name = 'Arial'
    r_opt0.font.size = Pt(8)

    p_opt1 = opt_table.rows[0].cells[1].paragraphs[0]
    p_opt1.paragraph_format.space_after = Pt(0)
    r_opt1 = p_opt1.add_run("[   ] Excellent       [   ] Good       [   ] Needs Minor Changes       [   ] Needs Major Changes")
    r_opt1.font.name = 'Arial'
    r_opt1.font.size = Pt(8)
    r_opt1.font.color.rgb = RGBColor(75, 85, 99)

    p_opt2 = opt_table.rows[1].cells[0].paragraphs[0]
    p_opt2.paragraph_format.space_after = Pt(0)
    r_opt2 = p_opt2.add_run("Would you approve this version for launch?")
    r_opt2.bold = True
    r_opt2.font.name = 'Arial'
    r_opt2.font.size = Pt(8)

    p_opt3 = opt_table.rows[1].cells[1].paragraphs[0]
    p_opt3.paragraph_format.space_after = Pt(0)
    r_opt3 = p_opt3.add_run("[   ] Yes       [   ] No")
    r_opt3.font.name = 'Arial'
    r_opt3.font.size = Pt(8)
    r_opt3.font.color.rgb = RGBColor(75, 85, 99)

    p_spacer_ap2 = doc.add_paragraph()
    p_spacer_ap2.paragraph_format.space_before = Pt(3)
    p_spacer_ap2.paragraph_format.space_after = Pt(3)

    # Feedback comments box
    feed_table = doc.add_table(rows=1, cols=1)
    feed_table.autofit = False
    feed_cell = feed_table.cell(0, 0)
    set_cell_background(feed_cell, "FFFDF8")
    feed_cell.width = Inches(7.5)
    set_cell_margins(feed_cell, top=80, bottom=600, left=100, right=100)
    
    f_tcPr = feed_cell._tc.get_or_add_tcPr()
    f_tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E4D8C8"/>\n'
        f'</w:tcBorders>'
    )
    f_tcPr.append(f_tcBorders)

    fp = feed_cell.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    frun = fp.add_run("ADDITIONAL CLIENT FEEDBACK:")
    frun.font.name = 'Arial'
    frun.font.size = Pt(7.5)
    frun.bold = True
    frun.font.color.rgb = RGBColor(75, 85, 99)

    p_spacer_ap3 = doc.add_paragraph()
    p_spacer_ap3.paragraph_format.space_before = Pt(6)
    p_spacer_ap3.paragraph_format.space_after = Pt(6)

    # Signature blocks
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    sig_table.rows[0].cells[0].width = Inches(4.7)
    sig_table.rows[0].cells[1].width = Inches(2.8)
    sig_table.rows[1].cells[0].width = Inches(4.7)
    sig_table.rows[1].cells[1].width = Inches(2.8)

    for r in sig_table.rows:
        for c in r.cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    p_s0 = sig_table.rows[0].cells[0].paragraphs[0]
    p_s0.paragraph_format.space_after = Pt(0)
    r_s0 = p_s0.add_run("Client Name (Printed): _________________________________________")
    r_s0.font.name = 'Arial'
    r_s0.font.size = Pt(8)

    p_s1 = sig_table.rows[0].cells[1].paragraphs[0]
    p_s1.paragraph_format.space_after = Pt(0)
    r_s1 = p_s1.add_run("Date: _________________________")
    r_s1.font.name = 'Arial'
    r_s1.font.size = Pt(8)

    p_s2 = sig_table.rows[1].cells[0].paragraphs[0]
    p_s2.paragraph_format.space_after = Pt(0)
    r_s2 = p_s2.add_run("Signature: __________________________________________________")
    r_s2.font.name = 'Arial'
    r_s2.font.size = Pt(8)

    # Page Break for Dashboard Summary
    doc.add_page_break()

    # Final Dashboard Summary Title
    p_d_title = doc.add_paragraph()
    r_d_title = p_d_title.add_run("INTIKHAB PLATFORM VERIFICATION SUMMARY")
    r_d_title.font.name = 'Arial'
    r_d_title.font.size = Pt(18)
    r_d_title.bold = True
    r_d_title.font.color.rgb = RGBColor(29, 26, 23)
    p_d_title.paragraph_format.space_after = Pt(2)

    p_d_sub = doc.add_paragraph()
    r_d_sub = p_d_sub.add_run("Launch Readiness Dashboard Status")
    r_d_sub.font.name = 'Arial'
    r_d_sub.font.size = Pt(11)
    r_d_sub.bold = True
    r_d_sub.font.color.rgb = RGBColor(180, 35, 24)
    p_d_sub.paragraph_format.space_after = Pt(10)

    # Summary Dashboard Table (4 columns)
    dash_table = doc.add_table(rows=9, cols=4)
    dash_table.autofit = False
    style_table_borders(dash_table)
    
    col_widths_dash = [Inches(4.3), Inches(1.0), Inches(1.2), Inches(1.0)]

    # Header Row
    headers_dash = ["Platform Module Component", "Passed", "Needs Changes", "Not Tested"]
    for c_idx, h_text in enumerate(headers_dash):
        cell = dash_table.rows[0].cells[c_idx]
        cell.width = col_widths_dash[c_idx]
        set_cell_background(cell, "F6EFE5")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if c_idx > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(h_text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(29, 26, 23)

    modules_list = [
        "Homepage",
        "Products",
        "Cart",
        "Checkout",
        "Account",
        "Admin",
        "Mobile",
        "Performance"
    ]

    for idx, mod in enumerate(modules_list, 1):
        row = dash_table.rows[idx]
        bg_color = "FFFDF8" if idx % 2 == 1 else "FFFFFF"
        
        for c_idx in range(4):
            cell = row.cells[c_idx]
            cell.width = col_widths_dash[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)

            if c_idx == 0:
                run = p.add_run(mod)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(29, 26, 23)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("[   ]")
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(156, 163, 175)

    p_spacer_d = doc.add_paragraph()
    p_spacer_d.paragraph_format.space_before = Pt(14)
    p_spacer_d.paragraph_format.space_after = Pt(14)

    p_f = doc.add_paragraph()
    run_f = p_f.add_run("Intikhab E-Commerce Platform Client Checklist — Confidential Review Copy. Generated via automated testing suites.")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(7.5)
    run_f.italic = True
    run_f.font.color.rgb = RGBColor(107, 114, 128)

    doc.save(docx_filename)
    print("DOCX generated successfully.")

if __name__ == "__main__":
    generate_pdf()
    try:
        import docx
        generate_docx()
    except ImportError:
        print("python-docx is not installed yet. Skipping DOCX generation.")
        sys.exit(1)
